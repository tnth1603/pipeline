# ── IMPORTS ───────────────────────────────────────────────────────────────────
from flask import Flask, request, jsonify
import anthropic, json, os, base64, time, re
import pandas as pd
import numpy as np
from io import StringIO, BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches

# ── APP SETUP ─────────────────────────────────────────────────────────────────
app = Flask(__name__)
client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])

# ── CONFIG ────────────────────────────────────────────────────────────────────
COMPANY_NAME   = "Global Monkeypox Tracker"
MAX_ROWS       = 1000
MAX_ITERATIONS = 1
MAX_RETRIES    = 3

# ─────────────────────────────────────────────────────────────────────────────
# HELPER: Claude call with retry
# ─────────────────────────────────────────────────────────────────────────────
def call_claude(messages, max_tokens=1000, system=None):
    for attempt in range(MAX_RETRIES):
        try:
            kwargs = {
                "model": "claude-sonnet-4-20250514",
                "max_tokens": max_tokens,
                "messages": messages
            }
            if system:
                kwargs["system"] = system
            msg = client.messages.create(**kwargs)
            return msg.content[0].text
        except Exception as e:
            if attempt == MAX_RETRIES - 1:
                raise RuntimeError(f"Claude API failed after {MAX_RETRIES} attempts: {str(e)}")
            time.sleep(2 ** attempt)

# ─────────────────────────────────────────────────────────────────────────────
# HELPER: Parse JSON from Claude safely
# FIX: handles empty responses, markdown fences, and trailing text
# ─────────────────────────────────────────────────────────────────────────────
def parse_claude_json(raw):
    """
    Robustly parse JSON from Claude response.
    Handles: ```json fences, trailing text, empty responses.
    """
    if not raw or not raw.strip():
        raise ValueError("Claude returned empty response")

    # Strip markdown fences
    cleaned = raw.strip()
    cleaned = re.sub(r'^```json\s*', '', cleaned)
    cleaned = re.sub(r'^```\s*', '', cleaned)
    cleaned = re.sub(r'\s*```$', '', cleaned)
    cleaned = cleaned.strip()

    # Try direct parse first
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    # Try to extract first JSON object from response
    match = re.search(r'\{.*\}', cleaned, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            pass

    # Try to extract first JSON array
    match = re.search(r'\[.*\]', cleaned, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            pass

    raise ValueError(f"Could not parse JSON from Claude response: {cleaned[:200]}")

# ─────────────────────────────────────────────────────────────────────────────
# HELPER: Pre-compute global stats (domain-agnostic)
# ─────────────────────────────────────────────────────────────────────────────
def compute_global_stats(df_full):
    stats = {
        "total_rows":    int(len(df_full)),
        "total_columns": int(len(df_full.columns)),
        "column_names":  list(df_full.columns),
        "date_range":    "N/A",
        "unique_locations": 0,
    }

    for date_col in ['date', 'Date', 'DATE', 'timestamp', 'period']:
        if date_col in df_full.columns:
            try:
                stats["date_range"] = f"{df_full[date_col].min()} to {df_full[date_col].max()}"
            except Exception:
                pass
            break

    for loc_col in ['location', 'country', 'region', 'category', 'Location', 'Country']:
        if loc_col in df_full.columns:
            try:
                stats["unique_locations"] = int(df_full[loc_col].nunique())
                stats["location_column"]  = loc_col
            except Exception:
                pass
            break

    numeric_cols = df_full.select_dtypes(include='number').columns.tolist()
    stats["numeric_summaries"] = {}
    for col in numeric_cols[:8]:
        try:
            stats["numeric_summaries"][col] = {
                "sum":   round(float(df_full[col].sum()), 2),
                "max":   round(float(df_full[col].max()), 2),
                "mean":  round(float(df_full[col].mean()), 2),
                "nulls": int(df_full[col].isnull().sum())
            }
        except Exception:
            pass

    try:
        loc_col = stats.get("location_column", None)
        if loc_col and numeric_cols:
            first_num = numeric_cols[0]
            if 'iso_code' in df_full.columns:
                countries = df_full[~df_full['iso_code'].str.startswith('OWID', na=False)]
            else:
                countries = df_full
            top10 = (
                countries.groupby(loc_col)[first_num]
                .max().nlargest(10).reset_index()
                .to_dict(orient='records')
            )
            stats["top_10_items"]    = top10
            stats["top_items_metric"] = first_num
    except Exception:
        stats["top_10_items"] = []

    try:
        if 'total_cases' in df_full.columns and 'total_deaths' in df_full.columns:
            world = df_full[df_full['location'].str.contains('World|OWID_WRL', na=False)]
            if not world.empty:
                tc = int(world['total_cases'].max())
                td = int(world['total_deaths'].max())
                stats["total_global_cases"]     = tc
                stats["total_global_deaths"]    = td
                stats["case_fatality_rate_pct"] = round((td / tc * 100), 4) if tc > 0 else 0.0
                peak_row = world.loc[world['new_cases'].idxmax()]
                stats["peak_date"]        = str(peak_row['date'])
                stats["peak_daily_cases"] = int(peak_row['new_cases'])
    except Exception:
        pass

    return stats

# ─────────────────────────────────────────────────────────────────────────────
# HELPER: Stratified sample
# ─────────────────────────────────────────────────────────────────────────────
def stratified_sample(df_full, max_rows=MAX_ROWS):
    if len(df_full) <= max_rows:
        return df_full
    for loc_col in ['location', 'country', 'region', 'category']:
        if loc_col in df_full.columns:
            try:
                sampled = (
                    df_full.groupby(loc_col, group_keys=False)
                    .apply(lambda x: x.sample(
                        max(1, int(max_rows * len(x) / len(df_full))),
                        random_state=42
                    ))
                )
                return sampled.head(max_rows)
            except Exception:
                pass
    return df_full.sample(n=max_rows, random_state=42)

# ─────────────────────────────────────────────────────────────────────────────
# STEP 2: AGENT CLEANING
# ─────────────────────────────────────────────────────────────────────────────
CLEANING_AGENT_PROMPT = """
You are an expert data cleaning agent. You will receive a statistical summary of a dataset.
Identify ONE data quality issue that has not yet been fixed.
Return ONLY a JSON object with NO other text before or after it:

{
  "issue_found": true,
  "issue_type": "null_values | duplicates | outlier | format_inconsistency | placeholder",
  "column": "column name affected, or 'multiple'",
  "description": "Plain English description of the issue",
  "reasoning": "Why this is an error, not valid data",
  "confidence": "high | medium | low",
  "action": {
    "type": "one of: fillna_median | fillna_zero | fillna_mode | drop_duplicates | drop_nulls | replace_value | normalize_case | to_datetime | clip_outliers",
    "column": "column name to apply action to",
    "value": "only for replace_value: the value to replace",
    "replacement": "only for replace_value: what to replace it with"
  },
  "needs_human_review": false,
  "human_review_reason": ""
}

If no issues remain return ONLY: { "issue_found": false }

CRITICAL: Return ONLY the JSON object. No explanation. No markdown. No text before or after.
"""

SAFE_ACTIONS = {
    "fillna_median": lambda df, col, **_: df.assign(
        **{col: df[col].fillna(df[col].median())}
    ) if col in df.columns and pd.api.types.is_numeric_dtype(df[col]) else df,
    "fillna_zero": lambda df, col, **_: df.assign(
        **{col: df[col].fillna(0)}
    ) if col in df.columns else df,
    "fillna_mode": lambda df, col, **_: df.assign(
        **{col: df[col].fillna(df[col].mode()[0])}
    ) if col in df.columns and not df[col].mode().empty else df,
    "drop_duplicates": lambda df, **_: df.drop_duplicates(),
    "drop_nulls": lambda df, col, **_: df.dropna(subset=[col]) if col in df.columns else df,
    "replace_value": lambda df, col, value, replacement, **_: df.assign(
        **{col: df[col].replace(value, replacement)}
    ) if col in df.columns else df,
    "normalize_case": lambda df, col, **_: df.assign(
        **{col: df[col].str.strip().str.title()}
    ) if col in df.columns and df[col].dtype == object else df,
    "to_datetime": lambda df, col, **_: df.assign(
        **{col: pd.to_datetime(df[col], errors='coerce')}
    ) if col in df.columns else df,
    "clip_outliers": lambda df, col, **_: df.assign(
        **{col: df[col].clip(lower=df[col].quantile(0.01), upper=df[col].quantile(0.99))}
    ) if col in df.columns and pd.api.types.is_numeric_dtype(df[col]) else df,
}

def safe_exec_fix(df, action):
    action_type = action.get("type", "")
    col         = action.get("column", "")
    value       = action.get("value", None)
    replacement = action.get("replacement", None)
    if action_type not in SAFE_ACTIONS:
        raise ValueError(f"Unknown action: '{action_type}'")
    return SAFE_ACTIONS[action_type](df, col=col, value=value, replacement=replacement)

def summarise_df(df, fixed_so_far):
    summary = {
        "shape": {"rows": int(df.shape[0]), "cols": int(df.shape[1])},
        "columns": {},
        "sample_rows": df.head(3).to_dict(orient='records'),
        "issues_already_fixed": [x.get('description', '') for x in fixed_so_far]
    }
    for col in df.columns:
        col_info = {
            "dtype":  str(df[col].dtype),
            "nulls":  int(df[col].isnull().sum()),
            "unique": int(df[col].nunique()),
            "sample": [str(x) for x in df[col].dropna().unique()[:5]]
        }
        if pd.api.types.is_numeric_dtype(df[col]) and not df[col].isnull().all():
            col_info["min"]  = float(df[col].min())
            col_info["max"]  = float(df[col].max())
            col_info["mean"] = round(float(df[col].mean()), 4)
        summary["columns"][col] = col_info
    return json.dumps(summary, indent=2, default=str)[:4000]

def agent_clean_csv(df_sample, max_iterations=MAX_ITERATIONS):
    df           = df_sample.copy()
    rows_before  = len(df)
    cleaning_log = []
    flags        = []

    for iteration in range(1, max_iterations + 1):
        try:
            summary  = summarise_df(df, cleaning_log)
            raw      = call_claude(
                system=CLEANING_AGENT_PROMPT,
                messages=[{"role": "user", "content": f"Inspect this dataset summary:\n{summary}"}],
                max_tokens=500
            )
            # FIX: use robust JSON parser instead of raw json.loads
            decision = parse_claude_json(raw)
        except Exception as e:
            cleaning_log.append({"iteration": iteration, "action": f"error: {str(e)}"})
            break

        if not decision.get("issue_found", False):
            break

        action       = decision.get("action", {})
        description  = decision.get("description", "")
        confidence   = decision.get("confidence", "low")
        needs_review = decision.get("needs_human_review", False)

        if needs_review or confidence == "low" or not action:
            entry = {"iteration": iteration, "description": description, "action": "flagged for human review"}
            flags.append(entry)
            cleaning_log.append(entry)
            continue

        try:
            rows_before_fix = len(df)
            df = safe_exec_fix(df, action)
            cleaning_log.append({
                "iteration":    iteration,
                "issue_type":   decision.get("issue_type"),
                "description":  description,
                "action_taken": action,
                "rows_before":  rows_before_fix,
                "rows_after":   len(df),
                "action":       "applied"
            })
        except Exception as e:
            cleaning_log.append({"iteration": iteration, "description": description, "action": f"fix_failed: {str(e)}"})

    return {
        "clean_df":        df,
        "cleaning_log":    cleaning_log,
        "flags_for_human": flags,
        "rows_before":     rows_before,
        "rows_after":      len(df)
    }

# ─────────────────────────────────────────────────────────────────────────────
# STEP 3: ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
ANALYSIS_SYSTEM_PROMPT = """
You are a senior data analyst. You will receive:
1. Pre-computed accurate statistics from the FULL dataset
2. A sample of raw rows for pattern analysis

Use the pre-computed stats for all KPI values.
Use the raw rows only to identify trends and patterns.

Return ONLY a valid JSON object with NO other text:
{
  "period": "date range covered",
  "row_count": 0,
  "kpis": {"metric_name": value},
  "trends": ["trend 1", "trend 2", "trend 3"],
  "anomalies": ["anomaly 1"],
  "top_items": [{"name": "item", "value": 0, "label": "metric"}],
  "recommendations": ["recommendation 1", "recommendation 2", "recommendation 3"]
}
"""

def analyse_with_claude(clean_csv, global_stats):
    prompt = f"""
ACCURATE GLOBAL STATISTICS (from full dataset of {global_stats.get('total_rows','?')} rows):
{json.dumps(global_stats, indent=2)}

SAMPLE ROWS FOR PATTERN ANALYSIS:
{clean_csv[:5000]}

Use the global statistics for all KPI values. Infer domain from column names.
"""
    raw = call_claude(
        system=ANALYSIS_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1000
    )
    return parse_claude_json(raw)

# ─────────────────────────────────────────────────────────────────────────────
# STEP 4: DASHBOARD
# ─────────────────────────────────────────────────────────────────────────────
def generate_dashboard(insights, company=COMPANY_NAME):
    prompt = f"""
Create a complete self-contained HTML dashboard for {company}.
Use Chart.js from https://cdn.jsdelivr.net/npm/chart.js
Include:
- Header: {company} | Period: {insights.get('period','')}
- KPI cards: {json.dumps(insights.get('kpis',{}))}
- Bar chart of top items: {json.dumps(insights.get('top_items',[])[:8])}
- Trends section and Recommendations section
Dark professional theme (#0f172a background, white text, teal accents #4ecdc4).
Return ONLY complete HTML. No markdown. No backticks.
"""
    html = call_claude(messages=[{"role": "user", "content": prompt}], max_tokens=3000)
    return re.sub(r'^```html\s*|^```\s*|\s*```$', '', html.strip()).strip()

# ─────────────────────────────────────────────────────────────────────────────
# STEP 5: WORD REPORT
# FIX 1: parse_inline_bold() strips **markdown** and renders real Word bold
# FIX 2: cleaning log error shows friendly message instead of raw JSON error
# ─────────────────────────────────────────────────────────────────────────────
REPORT_PROMPT = """
Write a professional analytics report based on these insights.
Use these exact section headings (prefix each with #):
# Executive Summary
# Key Performance Indicators
# Trend Analysis
# Anomalies & Risk Signals
# Strategic Recommendations

IMPORTANT: Do NOT use **bold** markdown syntax. Write in plain text only.
Use plain sentences. No asterisks. No markdown formatting of any kind.
Concise, evidence-based, written for senior managers.
Insights: {insights}
"""

def write_report(insights):
    return call_claude(
        messages=[{"role": "user", "content": REPORT_PROMPT.format(insights=json.dumps(insights, indent=2))}],
        max_tokens=1500
    )

def parse_inline_bold(paragraph, text):
    """
    Parse **bold** markdown and add properly formatted runs to a Word paragraph.
    Handles the case where Claude ignores the no-markdown instruction.
    """
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.size = Pt(11)
        if i % 2 == 1:  # odd index = content inside **
            run.bold = True

def make_docx(report_text, insights, cleaning_log=None):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Title
    title = doc.add_heading(f"{COMPANY_NAME} — Analytics Report", 0)
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
        title.runs[0].font.size = Pt(20)

    sub = doc.add_paragraph(
        f"Period: {insights.get('period','N/A')} | "
        f"Total rows analysed: {insights.get('row_count','N/A')}"
    )
    if sub.runs:
        sub.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
        sub.runs[0].font.size = Pt(10)
    doc.add_paragraph("")

    # KPI table
    doc.add_heading("KPI Snapshot", 1)
    kpis = insights.get("kpis", {})
    if kpis:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Value"
        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        for k, v in kpis.items():
            row = table.add_row().cells
            row[0].text = str(k).replace("_", " ").title()
            row[1].text = str(v)
            for cell in row:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
    doc.add_paragraph("")

    # Report body
    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        elif line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), 1)
            if h.runs:
                h.runs[0].font.size = Pt(14)
        elif line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), 2)
            if h.runs:
                h.runs[0].font.size = Pt(12)
        elif line.startswith(("- ", "* ", "• ")):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_bold(p, line[2:].strip())
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            parse_inline_bold(p, re.sub(r'^\d+\.\s', '', line))
        else:
            p = doc.add_paragraph()
            parse_inline_bold(p, line)

    

    buf = BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode()

# ─────────────────────────────────────────────────────────────────────────────
# /run ENDPOINT
# ─────────────────────────────────────────────────────────────────────────────
@app.route("/run", methods=["POST"])
def run_pipeline():
    token = request.headers.get("X-Pipeline-Secret", "")
    if token != os.environ.get("PIPELINE_SECRET", ""):
        return jsonify({"error": "Unauthorized"}), 401

    try:
        data     = request.get_json(force=True, silent=True) or {}
        csv_text = data.get("csv", "")
        company  = data.get("company", COMPANY_NAME)

        if not csv_text:
            return jsonify({"error": "No CSV provided"}), 400

        try:
            df_full = pd.read_csv(StringIO(csv_text))
        except Exception as e:
            return jsonify({"error": f"CSV parse failed: {str(e)}"}), 400

        global_stats    = compute_global_stats(df_full)
        df_sample       = stratified_sample(df_full, max_rows=MAX_ROWS)
        cleaning_result = agent_clean_csv(df_sample, max_iterations=MAX_ITERATIONS)
        df_clean        = cleaning_result["clean_df"]
        clean_csv_text  = df_clean.to_csv(index=False)
        cleaning_log    = cleaning_result["cleaning_log"]

        insights       = analyse_with_claude(clean_csv_text, global_stats)
        dashboard_html = generate_dashboard(insights, company)
        report_text    = write_report(insights)
        docx_b64       = make_docx(report_text, insights, cleaning_log=cleaning_log)

        date_str = pd.Timestamp.today().strftime("%Y-%m-%d")

        return jsonify({
            "docx_b64":        docx_b64,
            "docx_filename":   f"Report_{date_str}.docx",
            "dashboard_html":  dashboard_html,
            "dash_filename":   f"Dashboard_{date_str}.html",
            "insights":        insights,
            "global_stats":    global_stats,
            "cleaning_log":    cleaning_log,
            "flags_for_human": cleaning_result["flags_for_human"],
            "cleaning_stats": {
                "rows_in_full_dataset": len(df_full),
                "rows_sampled":         len(df_sample),
                "rows_after_cleaning":  cleaning_result["rows_after"]
            },
            "status": "success"
        })

    except Exception as e:
        return jsonify({"error": str(e), "status": "failed"}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8080)
